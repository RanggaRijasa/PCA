#!/usr/bin/env python3
"""Create harmless PDF, DOCX, TXT, and MD fixtures for local ingestion checks."""

from __future__ import annotations

import json
from pathlib import Path
import sys
from typing import Dict

import fitz
from docx import Document

PROJECT_ROOT = Path(__file__).resolve().parents[1]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from backend.config import settings


def _write_metadata(path: Path, title: str, department: str) -> None:
    metadata: Dict[str, object] = {
        "title": title,
        "document_type": "policy",
        "department": department,
        "access_level": "staff",
        "owner": f"{department} Team",
        "version": "1.0",
        "effective_date": "2026-01-01",
        "expiry_date": None,
    }
    sidecar = path.with_name(f"{path.name}.metadata.json")
    sidecar.write_text(
        json.dumps(metadata, indent=2, ensure_ascii=True),
        encoding="utf-8",
    )


def _create_markdown(raw_path: Path) -> Path:
    path = raw_path / "sample_remote_work.md"
    path.write_text(
        "# Purpose\n"
        "This dummy policy explains a fictional remote-work request process.\n\n"
        "## Request Process\n"
        "Staff submit a request to a manager before changing their work location.\n\n"
        "## Equipment\n"
        "Only approved test equipment may be used for this fictional example.\n",
        encoding="utf-8",
    )
    _write_metadata(path, "Sample Remote Work Policy", "People")
    return path


def _create_text(raw_path: Path) -> Path:
    path = raw_path / "sample_security_guide.txt"
    path.write_text(
        "DEVICE USE\n"
        "Use approved test devices and lock the screen when stepping away.\n\n"
        "INCIDENT REPORTING\n"
        "Report fictional security incidents to the sample help desk.\n",
        encoding="utf-8",
    )
    _write_metadata(path, "Sample Security Guide", "Security")
    return path


def _create_docx(raw_path: Path) -> Path:
    path = raw_path / "sample_leave_handbook.docx"
    document = Document()
    document.add_heading("Leave Requests", level=1)
    document.add_paragraph(
        "This dummy handbook says staff submit leave requests to a manager."
    )
    document.add_heading("Example Allowances", level=2)
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Leave type"
    table.cell(0, 1).text = "Example days"
    table.cell(1, 0).text = "Annual"
    table.cell(1, 1).text = "12"
    document.save(path)
    _write_metadata(path, "Sample Leave Handbook", "People")
    return path


def _create_pdf(raw_path: Path) -> Path:
    path = raw_path / "sample_expense_policy.pdf"
    document = fitz.open()
    page = document.new_page()
    page.insert_text(
        (72, 72),
        "EXPENSE SUBMISSION\n"
        "This dummy policy requires sample receipts for fictional expenses.\n\n"
        "APPROVAL\n"
        "A manager reviews each fictional expense before reimbursement.",
    )
    document.save(path)
    document.close()
    _write_metadata(path, "Sample Expense Policy", "Finance")
    return path


def main() -> int:
    settings.raw_docs_path.mkdir(parents=True, exist_ok=True)
    created = [
        _create_markdown(settings.raw_docs_path),
        _create_text(settings.raw_docs_path),
        _create_docx(settings.raw_docs_path),
        _create_pdf(settings.raw_docs_path),
    ]
    print(f"Created {len(created)} dummy documents in {settings.raw_docs_path}:")
    for path in created:
        print(f"- {path.name}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
