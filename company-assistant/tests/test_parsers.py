"""Dummy-document parser coverage for all Phase 5 file types."""

from pathlib import Path

import fitz
from docx import Document

from ingestion.parsers import parse_document


def test_txt_and_markdown_preserve_sections(tmp_path: Path) -> None:
    text_path = tmp_path / "policy.txt"
    text_path.write_text(
        "ELIGIBILITY\nEmployees may apply.\n\nAPPROVAL\nManagers approve requests.",
        encoding="utf-8",
    )
    markdown_path = tmp_path / "guide.md"
    markdown_path.write_text(
        "# Overview\nGeneral guidance.\n\n## Steps\n1. Submit a request.",
        encoding="utf-8",
    )

    text_sections = parse_document(text_path).sections
    markdown_sections = parse_document(markdown_path).sections

    assert [section.section for section in text_sections] == [
        "ELIGIBILITY",
        "APPROVAL",
    ]
    assert [section.section for section in markdown_sections] == [
        "Overview",
        "Steps",
    ]


def test_docx_preserves_heading_and_table(tmp_path: Path) -> None:
    path = tmp_path / "handbook.docx"
    document = Document()
    document.add_heading("Leave", level=1)
    document.add_paragraph("Employees receive annual leave.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Type"
    table.cell(0, 1).text = "Days"
    table.cell(1, 0).text = "Annual"
    table.cell(1, 1).text = "12"
    document.save(path)

    parsed = parse_document(path)

    assert parsed.sections[0].section == "Leave"
    assert "Employees receive annual leave." in parsed.sections[0].text
    assert "| Type | Days |" in parsed.sections[0].text


def test_pdf_preserves_page_number(tmp_path: Path) -> None:
    path = tmp_path / "manual.pdf"
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), "SECURITY\nUse approved devices only.")
    document.save(path)
    document.close()

    parsed = parse_document(path)

    assert parsed.sections[0].page == 1
    assert parsed.sections[0].section == "SECURITY"
    assert "approved devices" in parsed.sections[0].text
