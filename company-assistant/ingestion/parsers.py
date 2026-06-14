"""Structured local parsers for PDF, DOCX, TXT, and Markdown files."""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import re
from typing import Iterable, List, Optional, Sequence, Tuple

import fitz
from docx import Document


class DocumentParseError(RuntimeError):
    """Raised when a document cannot be parsed reliably."""


@dataclass(frozen=True)
class ParsedSection:
    text: str
    section: Optional[str] = None
    page: Optional[int] = None


@dataclass(frozen=True)
class ParsedDocument:
    sections: List[ParsedSection]
    requires_ocr: bool = False


def _looks_like_heading(line: str) -> bool:
    stripped = line.strip()
    if not stripped or len(stripped) > 120:
        return False
    if re.match(r"^#{1,6}\s+\S", stripped):
        return True
    letters = [character for character in stripped if character.isalpha()]
    return bool(letters) and stripped.upper() == stripped and len(stripped.split()) <= 12


def _heading_text(line: str) -> str:
    return re.sub(r"^#{1,6}\s+", "", line.strip()).strip()


def _split_sections(
    text: str,
    page: Optional[int] = None,
    default_section: Optional[str] = None,
) -> List[ParsedSection]:
    sections: List[ParsedSection] = []
    current_heading = default_section
    buffer: List[str] = []

    def flush() -> None:
        body = "\n".join(buffer).strip()
        if body:
            sections.append(
                ParsedSection(text=body, section=current_heading, page=page)
            )
        buffer.clear()

    for line in text.splitlines():
        if _looks_like_heading(line):
            flush()
            current_heading = _heading_text(line)
        else:
            buffer.append(line)
    flush()
    return sections


def _parse_pdf(path: Path) -> ParsedDocument:
    sections: List[ParsedSection] = []
    pages_with_text = 0
    try:
        with fitz.open(path) as document:
            for page_index, page in enumerate(document, start=1):
                text = page.get_text("text", sort=True).strip()
                if text:
                    pages_with_text += 1
                    sections.extend(
                        _split_sections(
                            text,
                            page=page_index,
                            default_section=f"Page {page_index}",
                        )
                    )
            if document.page_count and pages_with_text == 0:
                raise DocumentParseError(
                    "PDF contains no extractable text and requires OCR."
                )
            requires_ocr = (
                document.page_count > 0
                and pages_with_text / document.page_count < 0.5
            )
    except DocumentParseError:
        raise
    except Exception as exc:
        raise DocumentParseError(f"PDF parsing failed: {exc}") from exc
    return ParsedDocument(sections=sections, requires_ocr=requires_ocr)


def _table_to_markdown(table: object) -> str:
    rows = []
    for row in table.rows:
        rows.append([cell.text.strip().replace("\n", " ") for cell in row.cells])
    if not rows or not rows[0]:
        return ""
    width = max(len(row) for row in rows)
    normalized = [row + [""] * (width - len(row)) for row in rows]
    header = normalized[0]
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(["---"] * width) + " |",
    ]
    lines.extend("| " + " | ".join(row) + " |" for row in normalized[1:])
    return "\n".join(lines)


def _parse_docx(path: Path) -> ParsedDocument:
    try:
        document = Document(path)
    except Exception as exc:
        raise DocumentParseError(f"DOCX parsing failed: {exc}") from exc

    sections: List[ParsedSection] = []
    current_heading: Optional[str] = None
    buffer: List[str] = []

    def flush() -> None:
        body = "\n\n".join(part for part in buffer if part).strip()
        if body:
            sections.append(ParsedSection(text=body, section=current_heading))
        buffer.clear()

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if paragraph.style and paragraph.style.name.lower().startswith("heading"):
            flush()
            current_heading = text
        else:
            buffer.append(text)
    for table in document.tables:
        table_text = _table_to_markdown(table)
        if table_text:
            buffer.append(table_text)
    flush()

    if not sections:
        raise DocumentParseError("DOCX contains no extractable text.")
    return ParsedDocument(sections=sections)


def _parse_text(path: Path) -> ParsedDocument:
    try:
        text = path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        try:
            text = path.read_text(encoding="utf-8-sig")
        except UnicodeDecodeError as exc:
            raise DocumentParseError("Text file is not valid UTF-8.") from exc
    except OSError as exc:
        raise DocumentParseError(f"Text parsing failed: {exc}") from exc
    sections = _split_sections(text, default_section=path.stem)
    if not sections:
        raise DocumentParseError("Text file contains no extractable text.")
    return ParsedDocument(sections=sections)


def parse_document(path: Path) -> ParsedDocument:
    parsers = {
        ".pdf": _parse_pdf,
        ".docx": _parse_docx,
        ".txt": _parse_text,
        ".md": _parse_text,
    }
    parser = parsers.get(path.suffix.lower())
    if parser is None:
        raise DocumentParseError(f"No parser is available for {path.suffix}.")
    return parser(path)
